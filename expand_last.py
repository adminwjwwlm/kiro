# -*- coding: utf-8 -*-
"""Last expansion round - heavy content addition to both parts"""
from docx import Document
from docx.shared import Pt
from copy import deepcopy
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INPUT_FILE = '/projects/sandbox/kiro/技术方案/包6-国网新疆昌吉及阿克苏供电公司网络安全项目技术方案.docx'
OUTPUT_FILE = '/projects/sandbox/kiro/技术方案/包6-国网新疆昌吉及阿克苏供电公司网络安全项目技术方案.docx'

def insert_text_after_para(doc, para_idx, texts, font_size=152400):
    ref_para = doc.paragraphs[para_idx]
    for text in reversed(texts):
        new_para = OxmlElement('w:p')
        pPr_source = ref_para._element.find(qn('w:pPr'))
        if pPr_source is not None:
            new_pPr = deepcopy(pPr_source)
            numPr = new_pPr.find(qn('w:numPr'))
            if numPr is not None:
                new_pPr.remove(numPr)
            new_para.append(new_pPr)
        run_elem = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(int(font_size / 12700 * 2)))
        rPr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), str(int(font_size / 12700 * 2)))
        rPr.append(szCs)
        rn = OxmlElement('w:rFonts')
        rn.set(qn('w:eastAsia'), '仿宋_GB2312')
        rPr.append(rn)
        run_elem.append(rPr)
        t = OxmlElement('w:t')
        t.text = text
        t.set(qn('xml:space'), 'preserve')
        run_elem.append(t)
        new_para.append(run_elem)
        ref_para._element.addnext(new_para)

def find_para_index_by_text(doc, search_text, start_from=0):
    for i, para in enumerate(doc.paragraphs):
        if i < start_from:
            continue
        if search_text in para.text:
            return i
    return -1

print("Loading document for last expansion...")
doc = Document(INPUT_FILE)
print(f"Current: {len(doc.paragraphs)} paragraphs")

expansions = {}
part2_start = find_para_index_by_text(doc, "第二部分 国网新疆阿克苏供电公司")

# ==== PART 1 HEAVY ADDITIONS ====

# Add project organization and team description after 项目实施规模范围
idx = find_para_index_by_text(doc, "本项目的技术服务范围不包括变电站二次系统其他网络安全设备")
if idx > 0 and idx < part2_start:
    expansions[idx] = [
        "项目组织架构",
        "为确保项目的顺利实施，我们建立了科学合理的项目组织架构：",
        "（1）项目负责人：1名，全面负责项目的计划管理、进度控制、质量管理和客户沟通。项目负责人具有5年以上电力网络安全项目管理经验，熟悉国家电网公司网络安全管理体系和工作流程。",
        "（2）技术负责人：1名，负责升级方案的技术审核、技术难题的攻关解决、技术文档的质量把关。技术负责人具有网络安全高级工程师资质，精通各类防火墙设备的配置和管理。",
        "（3）升级操作人员：4名（分为A、B两组），负责现场升级操作的具体执行。所有操作人员经过专业培训和考核，持有相应的技术资格证书，熟悉本项目涉及的各型号防火墙设备的操作方法。",
        "（4）质量管控人员：1名，负责项目实施过程中的质量检查和监督，审核升级记录和验证报告的规范性和完整性。",
        "（5）后勤保障人员：1名，负责车辆调度、设备工具管理、介质安全管控等后勤支持工作。",
        "项目团队的人员配置充分考虑了19座变电站分布广泛的特点，通过两组操作人员的并行作业可以显著提高升级效率。同时，项目负责人、技术负责人和质量管控人员的设置保证了项目管理的规范性和技术方案的可靠性。",
        "项目团队各成员的职责分工明确、协作配合紧密。项目负责人统筹全局，技术负责人保障技术质量，操作人员执行具体任务，质量管控人员确保过程规范，后勤保障人员提供支持服务。各角色各司其职、相互配合，形成高效运转的项目团队。",
    ]

# Expand 国家法律法规 section detail
idx = find_para_index_by_text(doc, "国家标准和行业标准")
if idx > 0 and idx < part2_start:
    expansions[idx - 1] = [
        "以上法律法规构成了本项目实施的最上层法律依据。从法律合规的角度，电力企业作为关键信息基础设施运营者，有法定义务保持网络安全防护设施的有效性。防火墙特征库过旧导致的防护能力下降，可能被视为未尽到安全保护义务，从而面临法律风险。通过本项目确保特征库及时更新，是履行法律义务的具体行为。",
    ]

# Add more in 服务质量目标
idx = find_para_index_by_text(doc, "升级成功率目标")
if idx > 0 and idx < part2_start:
    expansions[idx + 4] = [
        "质量目标保障措施",
        "为实现上述服务质量目标，我们建立了以下质量保障措施体系：",
        "（1）制度保障：建立完善的质量管理制度，包括操作规程、检查标准、问题处理流程等，确保所有工作有章可循、有据可查。",
        "（2）人员保障：所有参与操作的技术人员经过严格的岗前培训和考核，考核合格后方可上岗操作。培训内容包括操作规程、设备操作方法、异常处置流程、安全注意事项等。",
        "（3）流程保障：建立升级操作的标准化流程，每个操作步骤都有明确的操作要求、检查标准和异常处理指引。通过标准化流程减少人为失误导致的质量问题。",
        "（4）工具保障：配备性能可靠的操作工具和测试验证工具，定期对工具进行检查和校验，确保工具的可用性和准确性。",
        "（5）监督保障：质量管控人员对操作过程和操作成果进行独立的质量检查和监督，发现问题及时纠正，确保最终交付质量达标。",
    ]

# Expand 应急预案 section in Part 1
idx = find_para_index_by_text(doc, "恶劣天气影响")
if idx > 0 and idx < part2_start:
    expansions[idx] = [
        "（6）通信中断处置：如果升级现场与项目管理组之间通信中断，升级人员按照预定的独立操作方案继续工作，不依赖远程指导。通信恢复后第一时间汇报工作进展和遇到的问题。",
        "（7）人员安全事故：如果升级人员在途中或现场发生安全事故，立即启动人员安全应急响应，优先保障人员安全，暂停升级工作。事故处理完毕后重新评估和安排后续工作。",
        "（8）升级包版本错误：如果到达现场后发现升级包版本不正确或不适用于该站设备，立即停止升级操作，返回后重新确认正确版本并更换升级包。不得使用不确定版本的升级包执行操作。",
        "应急处置的基本原则是：人员安全第一、设备安全第二、进度服从质量。在任何紧急情况下，都必须首先确保人员安全，其次确保设备不受损害，不得为赶进度而忽视安全和质量要求。",
        "应急联系机制：项目实施期间建立24小时应急联系通道。项目负责人手机保持24小时畅通；与昌吉供电公司信通部值班人员建立直通联系；与设备厂商技术支持建立快速响应通道。紧急情况下可同时启动多个联系通道，确保问题得到快速响应和解决。",
    ]

# ==== PART 2 HEAVY ADDITIONS ====

# Add project team description for Part 2
idx = find_para_index_by_text(doc, "项目团队在站内的工作严格遵守变电站安全管理规定", part2_start)
if idx > 0:
    expansions[idx] = [
        "项目组织与人员配置",
        "本项目组建专业的项目实施团队，人员配置如下：",
        "（1）项目负责人：1名，负责项目全面管理，统筹协调各项工作安排，对接阿克苏供电公司相关部门。",
        "（2）监测系统工程师：2名，负责设备接入安全监测系统、探针策略调优、二次辨识提升等与安全监测系统相关的工作。要求具有电力行业网络安全监测系统的建设和运维经验，熟悉电力行业专用通信协议。",
        "（3）安全评估工程师：1名，负责漏洞扫描、安全基线检查、风险评估等工作。要求持有CISP或同等资质证书，具有丰富的安全评估实践经验。",
        "（4）安全加固工程师：1名，负责设备安全加固、漏洞修复、端口封堵、代理策略整改等工作。要求具有多种类型设备的安全加固经验，熟悉各厂商设备的管理操作。",
        "（5）验证测试工程师：1名，负责各项工作完成后的效果验证和测试工作。要求具有系统性的测试验证能力，能够设计有效的验证方案和测试用例。",
        "（6）技术负责人：1名（可由监测系统高级工程师兼任），负责技术方案的审核把关、技术难题的攻关解决、技术文档的质量审核。",
        "项目团队人员的专业背景涵盖网络安全监测、漏洞评估、安全加固、系统测试等多个领域，能够全面覆盖项目各项工作的技术需求。团队成员均具有3年以上的电力行业网络安全从业经验，熟悉电力监控系统安全防护规定和相关技术标准，能够在严格遵守安全规定的前提下高效完成各项技术工作。",
    ]

# Add more in Part 2 探针策略 section
idx = find_para_index_by_text(doc, "告警聚合与降噪", part2_start)
if idx > 0:
    expansions[idx + 1] = [
        "探针策略精细化调优的具体技术方法详细说明如下：",
        "（1）基线分析方法：通过对正常网络流量的长时间观察和统计分析，建立每种设备类型、每个通信链路的流量基线模型。基线模型包括：通信时间分布特征（何时通信、通信频率）、流量大小特征（正常流量范围）、协议特征（使用哪些协议和端口）、行为特征（典型的操作序列）等多个维度。",
        "（2）规则分类管理：将探针检测规则按照威胁类型进行分类管理，包括：网络攻击类规则（如端口扫描检测、暴力破解检测）、恶意代码类规则（如木马通信特征检测）、异常行为类规则（如非工作时段访问、异常流量突增）、合规审计类规则（如违规外联检测、非法协议使用检测）等。各类规则独立管理，便于针对性优化。",
        "（3）阈值动态调整：对于基于阈值的检测规则（如连接频率阈值、流量大小阈值等），根据基线分析结果设定合理的阈值范围。阈值设置需要留有适当余量，既要避免正常波动触发误报，又要确保真实异常能够被及时检测。对于不同时段（如白天和夜间）可以设置差异化的阈值。",
        "（4）多条件关联检测：对于单一条件难以准确判断的场景，设计多条件关联检测规则。例如，单次登录失败可能是正常的输入错误，但短时间内大量登录失败则大概率是暴力破解攻击。通过设置时间窗口内的次数阈值实现关联检测。",
        "（5）策略效果评估方法：建立策略优化效果的定量评估体系。主要评估指标包括：误报率（FPR）、漏报率（FNR）、告警总量变化趋势、平均响应时间等。每次策略调整后统计这些指标的变化，评估调整效果，指导后续优化方向。",
    ]

# Add more detail in Part 2 安全加固
idx = find_para_index_by_text(doc, "日志审计加固", part2_start)
if idx > 0:
    expansions[idx + 1] = [
        "安全加固工作的实施流程和注意事项补充说明：",
        "实施流程遵循\"一设备一方案\"的原则，即为每台待加固设备制定独立的加固实施方案。方案内容包括：该设备当前不合规的配置项清单、每一项的具体加固操作步骤、操作可能产生的影响评估、回退操作步骤等。方案经技术负责人审核确认后执行。",
        "加固操作过程中需要特别注意以下事项：",
        "（1）操作顺序：先执行影响较小的加固项（如修改口令策略），后执行影响可能较大的加固项（如关闭服务）。这样即使后续操作出现问题，也可以在已完成的基础加固之上进行回退处理。",
        "（2）即时验证：每完成一个加固项后立即验证设备功能和网络通信是否正常。不要等到所有加固项全部完成后再统一验证，以免出现问题时难以定位是哪个加固操作导致的。",
        "（3）完整记录：详细记录每一步加固操作的执行时间、操作内容、操作结果和验证结果。记录信息要足够详细，能够在需要回退时准确还原每一步操作。",
        "（4）厂商协调：对于部分设备可能存在厂商特殊要求的加固项目（如某些配置项厂商建议不修改），事先与厂商确认加固方案的可行性，获得厂商的技术支持确认。",
    ]

# Add more in Part 2 future outlook - before the concluding paragraph  
idx = find_para_index_by_text(doc, "数字孪生与安全仿真", part2_start)
if idx > 0:
    expansions[idx + 1] = [
        "在技术发展路线方面，我们建议阿克苏供电公司制定变电站网络安全能力提升的中长期规划，分阶段、分步骤地实现从\"基本防护\"到\"深度防御\"再到\"智能安全\"的能力演进。本项目作为\"深度防御\"阶段的重要组成部分，将为后续向\"智能安全\"阶段的演进奠定坚实基础。",
        "从投资效益角度看，网络安全监测功能的提升是一项具有高投入产出比的投资。有效的安全监测能够在安全事件造成实际损失之前将其发现和阻断，避免的潜在损失远大于项目投入成本。同时，完善的安全监测体系也是企业网络安全能力成熟度的重要标志，有助于提升企业在行业内的安全管理水平和声誉。",
    ]

print(f"\nTotal insertion points: {len(expansions)}")
for para_idx in sorted(expansions.keys(), reverse=True):
    texts = expansions[para_idx]
    if para_idx >= len(doc.paragraphs):
        print(f"  WARNING: para_idx {para_idx} out of range, skipping")
        continue
    insert_text_after_para(doc, para_idx, texts)
    print(f"  Inserted {len(texts)} paragraphs after index {para_idx}")

print(f"\nTotal paragraphs: {len(doc.paragraphs)}")
print("Saving...")
doc.save(OUTPUT_FILE)

total_chars = sum(len(p.text) for p in doc.paragraphs)
# More realistic estimate: ~500 chars per page for 12pt Chinese with standard margins
est_pages_conservative = total_chars / 800
est_pages_realistic = total_chars / 550
print(f"Total characters: {total_chars}")
print(f"Estimated pages (conservative ~800 chars/page): {est_pages_conservative:.0f}")
print(f"Estimated pages (realistic ~550 chars/page): {est_pages_realistic:.0f}")
