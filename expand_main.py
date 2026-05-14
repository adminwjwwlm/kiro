#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
扩充包6技术方案文档
- 每部分扩充至约50页，整体100页
- 不更改文档结构和格式
- 2.1硬件软件实力章节保持空白
- 工作计划调整为2026年7月6日-12月31日
"""

from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from copy import deepcopy
import copy

INPUT_FILE = '/projects/sandbox/kiro/技术方案/包6-国网新疆昌吉及阿克苏供电公司网络安全项目技术方案.docx'
OUTPUT_FILE = '/projects/sandbox/kiro/技术方案/包6-国网新疆昌吉及阿克苏供电公司网络安全项目技术方案.docx'

def add_paragraph_after(doc, ref_para, text, font_size=Pt(12), bold=False, alignment=None):
    """在指定段落后插入新段落，复制格式"""
    new_para = deepcopy(ref_para._element)
    # Clear content
    from docx.oxml.ns import qn
    for child in list(new_para):
        if child.tag == qn('w:r'):
            new_para.remove(child)
    # Add new run
    from docx.oxml import OxmlElement
    run_elem = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(font_size.pt * 2)))
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), str(int(font_size.pt * 2)))
    rPr.append(szCs)
    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)
    run_elem.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    run_elem.append(t)
    new_para.append(run_elem)
    ref_para._element.addnext(new_para)
    return new_para


def insert_text_after_para(doc, para_index, texts, font_size=152400):
    """在指定段落索引后批量插入多段文字"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    ref_para = doc.paragraphs[para_index]
    # Insert in reverse order so they end up in correct order
    for text in reversed(texts):
        new_para = OxmlElement('w:p')
        # Copy paragraph properties from reference
        pPr_source = ref_para._element.find(qn('w:pPr'))
        if pPr_source is not None:
            new_pPr = deepcopy(pPr_source)
            # Remove numbering if any
            numPr = new_pPr.find(qn('w:numPr'))
            if numPr is not None:
                new_pPr.remove(numPr)
            new_para.append(new_pPr)
        
        # Add run with text
        run_elem = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(int(font_size / 12700 * 2)))  # Convert EMU to half-points
        rPr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), str(int(font_size / 12700 * 2)))
        rPr.append(szCs)
        rn = OxmlElement('w:rFonts')
        rn.set(qn('w:eastAsia'), '仿宋_GB2312')
        rPr.append(rn)
        run_elem.append(rPr)
        t = OxmlElement('w:t')
        t.text = text
        t.set(qn('xml:space'), 'preserve')
        run_elem.append(t)
        new_para.append(run_elem)
        
        ref_para._element.addnext(new_para)


def get_normal_font_size():
    """Return normal text font size in EMU (12pt = 152400)"""
    return 152400


print("Loading document...")
doc = Document(INPUT_FILE)
print(f"Loaded: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")

# We'll import and run the expansion modules
exec(open('/projects/sandbox/kiro/expand_content1.py', encoding='utf-8').read())
exec(open('/projects/sandbox/kiro/expand_content2.py', encoding='utf-8').read())
exec(open('/projects/sandbox/kiro/expand_schedule.py', encoding='utf-8').read())

print("\nSaving document...")
doc.save(OUTPUT_FILE)
print(f"Document saved to: {OUTPUT_FILE}")

# Check approximate page count
total_chars = sum(len(p.text) for p in doc.paragraphs)
# Approximate: ~800 chars per page for Chinese text in 12pt with margins
est_pages = total_chars / 800
print(f"Total characters: {total_chars}")
print(f"Estimated pages: {est_pages:.0f}")
print(f"Total paragraphs: {len(doc.paragraphs)}")
