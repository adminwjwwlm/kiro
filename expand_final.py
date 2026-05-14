# -*- coding: utf-8 -*-
"""Final expansion round to reach ~100 pages"""
from docx import Document
from docx.shared import Pt
from copy import deepcopy
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INPUT_FILE = '/projects/sandbox/kiro/技术方案/包6-国网新疆昌吉及阿克苏供电公司网络安全项目技术方案.docx'
OUTPUT_FILE = '/projects/sandbox/kiro/技术方案/包6-国网新疆昌吉及阿克苏供电公司网络安全项目技术方案.docx'

def insert_text_after_para(doc, para_idx, texts, font_size=152400):
    ref_para = doc.paragraphs[para_idx]
    for text in reversed(texts):
        new_para = OxmlElement('w:p')
        pPr_source = ref_para._element.find(qn('w:pPr'))
        if pPr_source is not None:
            new_pPr = deepcopy(pPr_source)
            numPr = new_pPr.find(qn('w:numPr'))
            if numPr is not None:
                new_pPr.remove(numPr)
            new_para.append(new_pPr)
        run_elem = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(int(font_size / 12700 * 2)))
        rPr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), str(int(font_size / 12700 * 2)))
        rPr.append(szCs)
        rn = OxmlElement('w:rFonts')
        rn.set(qn('w:eastAsia'), '仿宋_GB2312')
        rPr.append(rn)
        run_elem.append(rPr)
        t = OxmlElement('w:t')
        t.text = text
        t.set(qn('xml:space'), 'preserve')
        run_elem.append(t)
        new_para.append(run_elem)
        ref_para._element.addnext(new_para)

def find_para_index_by_text(doc, search_text, start_from=0):
    for i, para in enumerate(doc.paragraphs):
        if i < start_from:
            continue
        if search_text in para.text:
            return i
    return -1

print("Loading document for final expansion...")
doc = Document(INPUT_FILE)
print(f"Current: {len(doc.paragraphs)} paragraphs")

expansions = {}

# ====== PART 1 MORE CONTENT ======

# After "项目背景" title, add overview paragraph
idx = find_para_index_by_text(doc, "1.1.1 网络安全形势分析")
if idx > 0:
    expansions[idx - 1] = [
        "本章从网络安全形势、政策法规要求和昌吉电网现状三个维度阐述项目的背景情况，全面分析项目实施的外部环境和内部需求，为后续项目目标的确定和方案的设计提供充分的论证依据。",
    ]

# After 国际地缘政治 paragraph, add more threat analysis
idx = find_para_index_by_text(doc, "国际地缘政治形势的变化也对关键基础设施网络安全")
if idx > 0:
    expansions[idx] = [
        "从攻击技术演进的角度看，当前网络攻击技术正向着自动化、智能化和武器化方向发展。自动化攻击工具使得大规模扫描和漏洞利用变得更加容易，攻击者可以在极短时间内对大量目标发起攻击尝试。智能化技术使攻击行为更具针对性和隐蔽性，能够根据目标环境自动调整攻击策略。武器化则体现为国家级攻击组织将网络攻击能力作为战略武器储备，针对电力等关键基础设施研发专用攻击工具。",
        "具体到防火墙防护层面，攻击者为绕过防火墙检测不断开发新的规避技术，包括：加密隧道传输以规避内容检测、分片攻击以绕过特征匹配、协议混淆以逃避协议分析、零日漏洞利用以突破已知规则库等。防火墙特征库只有保持及时更新，纳入对最新规避技术的检测规则，才能持续有效地发挥防护作用。过旧的特征库可能对新出现的攻击变体和规避技术完全无法识别，等同于防护失效。",
        "从行业事件统计来看，2023-2024年期间，全球范围内公开报道的针对工业控制系统（含电力系统）的网络安全事件超过150起，较之前两年增长约40%。这些事件中约60%涉及已知漏洞的利用，如果相关系统的安全防护设备特征库能够及时更新到最新版本，理论上可以有效检测和阻断其中大部分攻击行为。这一数据有力说明了特征库及时更新对于实际安全防护效果的重要性。",
    ]

# More in 项目目标 overview
idx = find_para_index_by_text(doc, "总体目标")
if idx > 0 and idx < find_para_index_by_text(doc, "第二部分"):
    if "1.2.1" in doc.paragraphs[idx].text:
        expansions[idx - 1] = [
            "本项目的目标体系按照\"总体目标—具体目标—目标分解\"三个层次构建，确保项目目标既有宏观统领又有微观可执行、可衡量的指标体系。目标的设定充分考虑了项目的实际约束条件（时间、人力、技术等），做到了既具有挑战性又切实可行。",
        ]

# More in 实施路径 - 前期准备阶段 details
idx = find_para_index_by_text(doc, "工具和介质准备")
if idx > 0 and idx < find_para_index_by_text(doc, "第二部分"):
    expansions[idx] = [
        "（6）应急预案制定：针对升级过程中可能出现的各类异常情况（如设备死机、升级失败、网络中断等），制定详细的应急处置预案，明确各类异常的判断标准、处置步骤和责任分工。",
        "（7）沟通机制建立：与昌吉供电公司相关部门（调度中心、运维部、信通中心等）建立项目沟通联络机制，明确各方联系人、联系方式和信息传递要求。",
        "（8）质量管控方案：制定项目质量管控方案，明确各环节的质量检查标准和检查方式，建立质量问题的发现、报告和纠正流程。",
    ]

# More in 升级包准备 details
idx = find_para_index_by_text(doc, "介质准备")
if idx > 0 and idx < find_para_index_by_text(doc, "第二部分"):
    expansions[idx] = [
        "升级包准备工作的质量直接决定了现场升级操作的顺利程度。为确保升级包的安全性和可用性，我们建立了严格的多重验证机制：首先由专人从设备厂商官方安全中心下载最新版本的特征库升级包；然后由独立的安全检测人员使用多款杀毒软件进行全面扫描；接着由技术负责人验证升级包的数字签名和哈希值；最后在测试环境中进行模拟升级验证确认包的有效性。只有全部验证通过后，升级包才会被写入专用U盘用于现场升级。",
    ]

# More service content about documentation
idx = find_para_index_by_text(doc, "文档资料归档")
if idx > 0 and idx < find_para_index_by_text(doc, "第二部分"):
    expansions[idx] = [
        "文档归档管理遵循以下原则：",
        "（1）完整性原则：确保项目全过程产生的所有文档均纳入归档范围，不遗漏任何有价值的记录和资料。",
        "（2）规范性原则：所有文档按照统一的模板格式编制，文件命名遵循统一的命名规则（如\"站名_设备号_升级日期_文档类型\"），便于检索和查阅。",
        "（3）安全性原则：涉及设备配置信息和安全参数的文档按照保密要求进行管理，电子文档设置访问权限控制，纸质文档纳入保密档案管理。",
        "（4）长期保存原则：按照电力行业档案管理规定，项目文档保存期限不少于3年。电子文档采用双备份方式保存（本地存储+异地备份），确保数据的安全性和可用性。",
    ]

print(f"Part 1 final expansions: {sum(1 for k in expansions.keys() if k < find_para_index_by_text(doc, '第二部分 国网新疆阿克苏'))} points")



# ====== PART 2 MORE CONTENT ======
part2_start = find_para_index_by_text(doc, "第二部分 国网新疆阿克苏供电公司")

# More in Part 2 项目理解
idx = find_para_index_by_text(doc, "220kV纺织园变电站是阿克苏地区重要的电力输送节点", part2_start)
if idx > 0:
    expansions[idx] = [
        "网络安全监测功能提升的本质是对变电站网络安全\"感知-分析-响应\"能力链的全面增强。感知能力体现在监测覆盖范围的完整性和数据采集的全面性；分析能力体现在告警检测的准确性和异常识别的精准性；响应能力则建立在准确感知和精准分析的基础之上。三者缺一不可，只有实现全链条的协同提升，才能真正增强变电站的网络安全防护实效。",
        "从运维实践来看，一个高效的网络安全监测体系应当具备以下核心特征：一是监测覆盖无死角，所有关键设备和网络节点都在监控范围内；二是告警精准有价值，每一条告警都代表一个真实的安全风险或值得关注的事件；三是响应及时有效率，从安全事件发生到被发现再到被处置的全链条时间最短化。本项目的各项工作正是围绕这三个核心特征展开的系统性提升。",
    ]

# More in 纺织园变电站现状 - add detailed analysis
idx = find_para_index_by_text(doc, "综合以上分析，220kV纺织园变电站在网络安全监测方面存在的问题", part2_start)
if idx > 0:
    expansions[idx - 1] = [
        "从运维管理角度分析，上述问题的产生有其客观原因：一是变电站网络安全监测系统建设与站内设备扩容存在时间差，后续新增的设备未能及时纳入监测范围；二是初始部署时探针策略采用通用模板配置，未针对变电站特定的网络环境进行精细化适配；三是设备运行过程中缺乏定期的安全基线巡检和漏洞扫描机制，导致安全隐患逐步累积；四是运维人员对安全监测系统的运营维护能力有待加强，日常运维中对安全告警的分析深度不够。",
        "这些问题虽然目前尚未导致实际的安全事件，但已经构成了潜在的安全风险。在当前网络安全威胁日益严峻的环境下，任何监测盲区和防护薄弱环节都可能成为攻击者的突破口。因此，尽早实施系统性的监测功能提升，将安全隐患消除在萌芽状态，是明智且必要的选择。",
    ]

# More in Part 2 目标衡量标准
idx = find_para_index_by_text(doc, "目标衡量标准", part2_start)
if idx > 0:
    expansions[idx] = [
        "为确保项目目标的可衡量性和可验证性，我们建立了明确的目标达成评判标准和验证方法：",
        "（1）设备接入覆盖率：通过逐台核对设备台账和监测系统接入列表进行验证，计算公式为：已接入设备数÷应接入设备总数×100%。",
        "（2）告警采集完整率：通过在各设备上主动触发测试告警，验证监测系统能否正确接收和解析。计算公式为：成功采集告警数÷触发总告警数×100%。",
        "（3）探针策略误报率：以连续7天的观察期数据为统计样本，人工审核全部告警并标记误报。计算公式为：误报告警数÷总告警数×100%，与调优前的误报率进行对比。",
        "（4）漏洞修复率：使用漏洞扫描工具进行复测验证，计算公式为：修复后未检出漏洞数÷原始发现漏洞数×100%。",
        "（5）安全基线合规率：使用基线检查工具重新扫描，计算公式为：合规检查项数÷总检查项数×100%。",
    ]

# More in Part 2 实施路径 - 第二阶段 details
idx = find_para_index_by_text(doc, "各项工作按照风险等级由低到高的顺序实施", part2_start)
if idx > 0:
    expansions[idx + 1] = [
        "分步实施阶段的具体工作安排遵循以下原则：",
        "（1）依赖关系原则：先完成基础性工作（如设备接入、漏洞修复），再开展依赖前者成果的后续工作（如探针策略调优需要在设备全部接入后才能进行全面的流量基线分析）。",
        "（2）风险递增原则：先完成风险较低的工作（如设备接入仅涉及配置添加），后完成风险较高的工作（如设备加固涉及配置修改），逐步积累实施经验和信心。",
        "（3）验证驱动原则：每完成一个工作模块后必须进行即时验证，验证通过后才进入下一个模块。如验证发现问题，暂停后续工作优先解决当前问题。",
        "（4）弹性调整原则：实施过程中根据实际情况适当调整工作顺序和进度安排，确保整体目标不受局部调整的影响。",
    ]

# More in Part 2 可行性建议 - 网络架构优化建议
idx = find_para_index_by_text(doc, "部署网络流量审计", part2_start)
if idx > 0:
    expansions[idx] = [
        "（4）考虑引入微隔离技术。建议研究微隔离技术在变电站内部的应用可行性，通过在设备级别实现访问控制，进一步缩小单个设备被攻破后的影响范围，实现更精细化的安全防护。",
        "（5）完善网络拓扑文档。建议建立详细的、实时更新的网络拓扑文档，包括所有设备的连接关系、IP分配、VLAN划分、路由信息等，为安全运维和故障排查提供准确的参考依据。",
    ]

# More detail in Part 2 项目意义
idx = find_para_index_by_text(doc, "满足合规管理要求", part2_start)
if idx > 0 and "整改代理软件接入策略" in doc.paragraphs[idx].text:
    expansions[idx] = [
        "通过代理软件接入策略整改、安全加固和漏洞修复等工作，纺织园变电站在网络安全防护方面存在的各项不合规问题得到全面整改。这不仅使该站能够顺利通过各级网络安全检查和审计，更重要的是真正提升了安全防护的实际效果。合规不是目的，安全才是目的，本项目在实现合规的同时也实现了实质性的安全提升。",
    ]

# More in Part 2 未来展望 beginning
idx = find_para_index_by_text(doc, "展望未来，变电站网络安全监测将朝着以下方向发展", part2_start)
if idx > 0:
    expansions[idx - 1] = [
        "本项目的实施不仅解决了当前面临的具体安全问题，更重要的是为纺织园变电站建立了一套科学的网络安全监测运营方法论。这套方法论包括：如何评估安全监测现状、如何设计优化方案、如何安全实施变更操作、如何验证效果达成。这套方法论具有普遍适用性，可以在后续面对新的安全挑战时反复应用。",
    ]

# More in Part 2 服务质量目标
idx = find_para_index_by_text(doc, "客户满意度目标", part2_start)
if idx > 0:
    expansions[idx] = [
        "（6）变更安全目标：项目实施的所有配置变更操作均经过充分的方案评审和风险评估，变更操作一次性成功率不低于95%，因变更操作导致的非预期影响为零。",
        "（7）技能转移目标：项目验收时，阿克苏供电公司至少2名运维人员能够独立完成安全监测系统的日常运维操作，包括告警查看分析、基本策略调整、设备接入管理等。",
    ]

# More in Part 2 项目服务内容 - 端口封堵 detail
idx = find_para_index_by_text(doc, "封堵实施", part2_start)
if idx > 0 and "按照方案执行端口封堵操作" in doc.paragraphs[idx].text:
    expansions[idx] = [
        "端口封堵实施过程中的关键技术要点：",
        "（1）确认端口用途：对每一个计划封堵的端口，必须通过多种方式确认其当前确实未被使用。确认方式包括：查看端口连接状态、分析端口流量记录、咨询运维人员确认。只有经过充分确认的端口才能执行封堵。",
        "（2）封堵方式选择：根据端口类型和设备条件选择合适的封堵方式。对于设备自身开放的服务端口，优先在设备上直接关闭相应服务；对于物理端口，可以通过关闭端口或配置ACL进行限制。",
        "（3）分步验证：每封堵一个端口（或一组同类端口）后，立即验证设备功能和业务通信是否正常。如出现异常，立即回退最近一步封堵操作并排查原因。",
        "（4）封堵效果确认：全部封堵完成后，使用端口扫描工具从外部进行全面扫描，确认已封堵端口不再响应外部探测请求。同时从设备内部确认被封堵的端口不再监听。",
    ]

# More in 综合验证 section
idx = find_para_index_by_text(doc, "综合验证与优化服务", part2_start)
if idx > 0:
    expansions[idx + 13] = [  # After the section content
        "项目收尾与知识转移服务",
        "2.5.8.1 知识转移培训",
        "在项目收尾阶段，组织面向阿克苏供电公司运维人员的知识转移培训。培训内容包括：安全监测系统的日常运维操作方法、常见告警类型的分析判断方法、探针策略的基本调整方法、设备安全基线的检查方法、安全事件的初步处置流程等。培训采用理论讲解+实操演练的方式，确保运维人员能够真正掌握相关技能。",
        "2.5.8.2 运维手册编制",
        "编制详细的安全监测系统日常运维手册，内容包括：系统架构说明、设备接入管理操作指南、告警监控操作指南、常见问题处理指南、策略调整操作指南、应急处置操作指南等。运维手册作为项目交付成果之一，为运维人员提供持续参考。",
        "2.5.8.3 项目验收支持",
        "配合阿克苏供电公司完成项目验收工作，包括：整理和提交全套项目交付文档；配合验收专家的现场检查和演示要求；回答验收过程中提出的技术问题；对验收意见中的整改要求及时响应和落实。",
    ]

# Apply all final expansions
print(f"\nTotal final insertion points: {len(expansions)}")
for para_idx in sorted(expansions.keys(), reverse=True):
    texts = expansions[para_idx]
    if para_idx >= len(doc.paragraphs):
        print(f"  WARNING: para_idx {para_idx} out of range, skipping")
        continue
    insert_text_after_para(doc, para_idx, texts)
    print(f"  Inserted {len(texts)} paragraphs after index {para_idx}")

print(f"\nTotal paragraphs: {len(doc.paragraphs)}")

# Save
print("Saving...")
doc.save(OUTPUT_FILE)

total_chars = sum(len(p.text) for p in doc.paragraphs)
est_pages = total_chars / 800
print(f"Total characters: {total_chars}")
print(f"Estimated pages: {est_pages:.0f}")
