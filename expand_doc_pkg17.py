#!/usr/bin/env python3
"""Expand pkg17 technical proposal document to ~80000+ chars."""
import sys
sys.path.insert(0, '/projects/sandbox/kiro')
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

INPUT = '/projects/sandbox/kiro/技术方案/包17-国网新疆奎屯供电公司220kV、35kV变电站自动化设备网络安全探针接入服务项目及国网新疆奎屯供电公司110kV长春变等35座变电站自动化设备网络安全探针接入服务项目技术方案.docx'
OUTPUT = INPUT  # overwrite

doc = Document(INPUT)

def find_para_idx(text_fragment, start=0):
    """Find paragraph index containing text fragment."""
    for i, p in enumerate(doc.paragraphs):
        if i < start:
            continue
        if text_fragment in p.text:
            return i
    return -1

def insert_after(ref_para, text_content):
    """Insert a new paragraph after ref_para, copying its format."""
    new_p = OxmlElement('w:p')
    # Copy paragraph properties
    ref_pPr = ref_para._element.find(qn('w:pPr'))
    if ref_pPr is not None:
        new_pPr = deepcopy(ref_pPr)
        new_p.insert(0, new_pPr)
    # Create run
    new_r = OxmlElement('w:r')
    # Set run properties - 12pt fangsong
    rPr = OxmlElement('w:rPr')
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '24')
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), '24')
    rPr.append(szCs)
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rPr.append(rFonts)
    new_r.append(rPr)
    # Add text
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text_content
    new_r.append(t)
    new_p.append(new_r)
    ref_para._element.addnext(new_p)
    return new_p

def insert_multiple_after(ref_para, texts):
    """Insert multiple paragraphs after ref_para in order."""
    current = ref_para
    for text in texts:
        new_elem = insert_after_elem(current, text, ref_para)
        current_para = type('P', (), {'_element': new_elem})()
        current = current_para
    return current

def insert_after_elem(ref_obj, text_content, format_ref_para):
    """Insert after an element (para or wrapper), return new element."""
    new_p = OxmlElement('w:p')
    ref_pPr = format_ref_para._element.find(qn('w:pPr'))
    if ref_pPr is not None:
        new_pPr = deepcopy(ref_pPr)
        new_p.insert(0, new_pPr)
    new_r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '24')
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), '24')
    rPr.append(szCs)
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rPr.append(rFonts)
    new_r.append(rPr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text_content
    new_r.append(t)
    new_p.append(new_r)
    elem = ref_obj._element if hasattr(ref_obj, '_element') else ref_obj
    elem.addnext(new_p)
    return new_p

# Load content from external modules
exec(open('/projects/sandbox/kiro/pkg17_content.py').read())
exec(open('/projects/sandbox/kiro/pkg17_content2.py').read())
exec(open('/projects/sandbox/kiro/pkg17_content3.py').read())
exec(open('/projects/sandbox/kiro/pkg17_content4.py').read())
exec(open('/projects/sandbox/kiro/pkg17_content5.py').read())

print("Saving document...")
doc.save(OUTPUT)

# Verify
doc2 = Document(OUTPUT)
total = sum(len(p.text) for p in doc2.paragraphs)
print(f"Total paragraphs: {len(doc2.paragraphs)}")
print(f"Total characters: {total}")
if total >= 80000:
    print("SUCCESS: Document expanded to 80000+ characters")
else:
    print(f"WARNING: Only {total} chars, need more content")
