# -*- coding: utf-8 -*-
"""Expand pkg17 document to ~100 pages and adjust schedule to 2026.7.6-12.31"""
from docx import Document
from docx.shared import Pt
from copy import deepcopy
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INPUT_FILE = '/projects/sandbox/kiro/技术方案/包17-国网新疆奎屯供电公司220kV、35kV变电站自动化设备网络安全探针接入服务项目及国网新疆奎屯供电公司110kV长春变等35座变电站自动化设备网络安全探针接入服务项目技术方案.docx'

def insert_text_after_para(doc, para_idx, texts, font_size=152400):
    ref_para = doc.paragraphs[para_idx]
    for text in reversed(texts):
        new_para = OxmlElement('w:p')
        pPr_source = ref_para._element.find(qn('w:pPr'))
        if pPr_source is not None:
            new_pPr = deepcopy(pPr_source)
            numPr = new_pPr.find(qn('w:numPr'))
            if numPr is not None:
                new_pPr.remove(numPr)
            new_para.append(new_pPr)
        run_elem = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(int(font_size / 12700 * 2)))
        rPr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), str(int(font_size / 12700 * 2)))
        rPr.append(szCs)
        rn = OxmlElement('w:rFonts')
        rn.set(qn('w:eastAsia'), '仿宋_GB2312')
        rPr.append(rn)
        run_elem.append(rPr)
        t = OxmlElement('w:t')
        t.text = text
        t.set(qn('xml:space'), 'preserve')
        run_elem.append(t)
        new_para.append(run_elem)
        ref_para._element.addnext(new_para)

def find_para(doc, text, start=0):
    for i, p in enumerate(doc.paragraphs):
        if i < start: continue
        if text in p.text: return i
    return -1

print("Loading...")
doc = Document(INPUT_FILE)
print(f"Current: {len(doc.paragraphs)} paragraphs, {sum(len(p.text) for p in doc.paragraphs)} chars")

expansions = {}

# === PART 1 EXPANSIONS (220kV/35kV 64 stations) ===

# After project background intro
idx = find_para(doc, "国网新疆奎屯供电公司作为国网新疆电力有限公司的下属供电单位")
if idx > 0:
    expansions[idx] = [
        "从网络安全技术发展趋势来看，电力监控系统安全防护正从传统的边界防御模式向\"监测预警+主动防御\"的新模式转变。网络安全探针作为端点检测和响应（EDR）技术在电力行业的具体应用，通过在每台电力监控系统设备上部署轻量级安全代理，实现对设备运行状态、文件操作行为、网络通信行为、用户操作行为的实时监控和异常检测。这种基于主机的安全监测方式，能够有效弥补传统边界防护手段的不足，发现绕过边界防护的内部安全威胁。",
        "网络安全探针的技术优势体现在以下几个方面：一是监测视角更加全面，不仅能够监测网络层面的攻击行为，还能发现主机层面的异常操作（如关键文件篡改、非法程序运行、异常用户登录等）；二是检测精度更高，基于对每台设备正常行为基线的建模，能够精准识别偏离正常模式的异常行为；三是响应更加及时，安全事件从发生到被检测的时间可缩短到秒级；四是取证更加完整，能够记录安全事件的完整行为链，为事后分析和溯源提供有力支撑。",
        "对于奎屯供电公司辖区64座变电站而言，实现网络安全探针的全面部署和精细化配置，意味着从\"粗放式防护\"向\"精细化防护\"的重大转变，是建设智能化、自动化网络安全运营体系的关键基础设施。",
    ]

# Expand 项目重难点 - after "项目规模大、工期紧"
idx = find_para(doc, "如何在有限工期内高质量完成全部施工任务")
if idx > 0:
    expansions[idx] = [
        "（7）电压等级差异带来的标准化难题",
        "220kV和35kV变电站在设备配置、网络架构、安全防护等级等方面存在显著差异。220kV变电站通常配置15-25台设备，网络拓扑采用三层架构（站控层/间隔层/过程层），安全防护等级最高；而35kV变电站通常仅配置5-10台设备，网络架构较为简单。这种差异要求项目团队必须针对两种电压等级分别制定施工方案和策略模板，增加了方案设计的工作量和复杂性。",
        "应对措施：在试点阶段分别选取220kV和35kV变电站各1-2座进行验证，形成两套差异化的标准施工流程和策略配置模板。后续推进阶段，根据电压等级选用对应模板，在保证质量的前提下提高施工效率。",
        "（8）季节性气候因素影响",
        "奎屯地区位于新疆北部，冬季严寒漫长。如果项目执行期跨越冬季（11-12月），部分偏远地区35kV变电站的出行和现场作业将面临低温、冰雪等恶劣天气的影响。需要合理安排施工顺序，将偏远站点优先安排在天气条件较好的时段。",
        "应对措施：将偏远地区35kV变电站优先安排在7-9月施工；为冬季作业人员配备防寒保暖装备；遇极端天气及时调整计划，确保人员安全。",
    ]

# Expand 项目实施路径 - after Phase 3
idx = find_para(doc, "对不具备探针接入条件的设备及时列出清单报于现场负责人")
if idx > 0:
    expansions[idx] = [
        "（5）实施标准化质量管控。每座变电站施工完成后执行标准质量检查：探针安装完整性检查、策略配置合规性检查、告警触发有效性检查、系统运行稳定性检查。检查不合格的站点立即安排整改，整改合格后方可验收该站。",
        "（6）建立问题分类处理机制。将施工中遇到的问题分为三类：A类（可现场立即解决）、B类（需技术负责人协调解决）、C类（需设备厂商支持解决）。A类问题由施工人员自行处理；B类问题当日内解决；C类问题启动厂商支持流程，不影响其他站点施工进度。",
        "（7）进度跟踪与动态调整。项目经理每日更新施工进度看板，对比计划进度与实际进度。对于进度滞后的批次，分析原因并制定追赶措施（如增派人力、延长工时、调整批次安排等）。每周向甲方汇报进度情况。",
    ]

# Expand 可行性建议 after "(8)合理安排窗口期施工"
idx = find_para(doc, "利用设备停役的时机完成探针安装")
if idx > 0:
    expansions[idx] = [
        "（9）建立长效运维机制",
        "建议在项目完成后建立网络安全监测系统的长效运维机制。具体包括：制定安全监测系统日常巡检制度（每周至少一次系统状态检查）；建立策略动态优化机制（每月分析告警数据，持续优化策略配置）；建立探针版本管理制度（跟踪厂商发布的新版本，评估升级必要性）；建立人员培训计划（定期对运维人员进行安全监测系统操作培训）。",
        "（10）考虑与其他安全系统的联动",
        "建议在项目完成后，积极推动网络安全监测系统与其他安全防护系统的联动集成。例如：与防火墙联动实现自动封堵（当探针发现主机被入侵时，自动通知防火墙封锁相关IP）；与安全审计系统联动实现行为追溯（将探针告警与审计日志关联分析）；与态势感知平台联动实现全局态势展示。系统间的联动能够实现\"1+1>2\"的协同防御效果。",
        "（11）数据驱动的持续改进",
        "建议建立数据驱动的安全监测持续改进机制。通过对告警数据、设备运行数据、威胁情报数据的持续分析，发现安全防护的薄弱环节和改进机会。定期（每季度）编制安全态势分析报告，为管理层决策提供数据支撑。",
    ]

# Expand 项目意义
idx = find_para(doc, "为大规模项目实施提供示范")
if idx > 0:
    expansions[idx + 1] = [
        "（6）提升应急响应能力",
        "64座变电站全面部署网络安全探针后，一旦发生网络安全事件，安全运维人员能够通过集中管理平台快速定位受影响的设备范围、了解攻击路径和影响程度，为应急响应决策提供关键信息支撑。相比没有探针监测的情况，事件发现时间可从小时级缩短到分钟级，响应效率大幅提升。",
        "（7）形成安全资产管理基础",
        "项目实施过程中建立的完整资产台账和通信关系图谱，是网络安全管理的重要基础数据。这些数据不仅服务于本项目的探针部署和策略配置，还将为后续的风险评估、等级保护测评、安全审计等工作提供准确的资产底数和网络拓扑信息。",
    ]

# === PART 2 EXPANSIONS (110kV 35 stations) ===

# Find Part 2 start marker
p2_start = find_para(doc, "2、工作规划描述", 300)

# Expand Part 2 项目背景
idx = find_para(doc, "设备多样性带来的兼容性挑战", p2_start - 200)
if idx > 0:
    expansions[idx + 1] = [
        "针对设备兼容性问题，我们制定了系统化的评估和处理方案。在项目启动阶段，首先对35座变电站的设备清单进行全面梳理，按照设备品牌、型号、操作系统版本、投运年份等维度进行分类统计。对于每种设备类型，提前与设备厂商和探针软件厂商确认兼容性列表。对于不在兼容性列表内的设备，在试点阶段进行单独验证。对于经评估确实无法安装探针的老旧设备，及时列出清单报甲方确认，并提出替代监测方案（如通过网络流量镜像方式实现间接监测）。",
        "从操作系统分布来看，35座110kV变电站内的设备主要运行以下系统：Windows Server 2008/2012/2016/2019、Windows 7/10 Embedded、Linux（CentOS 6/7、Red Hat等）、以及各厂商自研的嵌入式操作系统。不同操作系统版本对应不同版本的探针软件，安装方式和配置方法各有差异。项目团队需要熟练掌握各种系统环境下的探针安装和配置技术。",
    ]

# Expand Part 2 项目实施路径 - Phase 3
idx = find_para(doc, "问题快速响应")
if idx > 0:
    expansions[idx + 2] = [
        "施工过程中的具体技术操作要点如下：",
        "（1）探针安装操作规范：安装前检查系统资源（CPU利用率<70%、可用内存>500MB、磁盘空间>2GB）；安装过程中监控系统性能指标变化；安装完成后验证探针服务正常启动且系统资源消耗在可接受范围内（探针进程CPU占用<5%、内存占用<200MB）。",
        "（2）策略配置操作规范：白名单配置前使用网络抓包工具采集设备通信关系（建议采集时间不少于4小时覆盖各时段通信）；配置完成后进入观察模式运行24小时，分析是否存在合法通信被拦截的情况；确认无误后切换为防护模式。",
        "（3）监测装置升级操作规范：升级前备份当前版本软件和配置文件；升级过程中保持另一台冗余装置在线运行；升级完成后验证双平面通信正常、数据上送正常、对时偏差<3秒；完成版本校验确认。",
        "（4）安全基线加固操作规范：加固前完整备份系统配置；逐项执行加固措施，每执行一项立即验证业务功能；加固后进行全面功能测试确认各项业务正常运行。",
    ]

# Expand Part 2 项目服务目标
idx = find_para(doc, "工期控制目标", p2_start - 100)
if idx > 0:
    expansions[idx] = [
        "严格按照合同约定工期完成全部施工任务。项目总工期为2026年7月6日至2026年12月31日。各阶段里程碑节点：准备阶段完成（2026年7月20日前）；试点阶段完成（2026年8月4日前）；全面推进阶段完成（2026年11月13日前）；验收收尾阶段完成（2026年12月31日前）。",
    ]

# Expand Part 2 项目服务内容 - 板块四 安全策略精细化配置
idx = find_para(doc, "板块四：安全策略精细化配置", p2_start - 50)
if idx > 0:
    expansions[idx] = [
        "安全策略精细化配置是本项目技术难度最大、工作量最重的核心工作。精细化配置的质量直接决定了安全监测系统的告警准确性和实际防护效果。以下对策略配置的主要内容进行详细说明：",
        "（1）关键目录配置：对Linux系统设备，配置监控的关键目录包括/bin、/sbin、/usr/bin、/usr/sbin、/etc、/boot、/lib、/usr/lib等系统关键路径下的文件变更；对Windows系统设备，配置监控Windows/System32、Windows/SysWOW64、Program Files等关键目录。任何对这些目录下文件的新增、修改、删除操作都将触发告警。",
        "（2）危险命令配置：对Linux系统配置监控的危险命令包括：系统启停类（shutdown、reboot、halt、init）、用户权限类（useradd、userdel、passwd、chmod 777）、网络操作类（ifconfig、route、iptables）、文件操作类（rm -rf、dd、mkfs）等。对Windows系统配置监控的危险命令包括：net user、net localgroup、sc、reg delete等。",
        "（3）网络连接白名单配置：基于每台设备实际的网络通信需求，精确配置允许的网络连接（包括源IP、目的IP、源端口、目的端口、协议类型）。非白名单内的网络连接将触发告警。白名单配置的准确性要求极高，配置不当会导致大量误报或合法通信被阻断。",
        "（4）服务端口白名单配置：明确每台设备正常运行所需开放的网络服务端口，非白名单端口的监听行为将触发告警。例如，远动装置通常只需开放IEC104通信端口（2404/2405）和SSH管理端口（22），其他端口的开放可能意味着异常。",
    ]

# Expand Part 2 建设依据 section
idx = find_para(doc, "电力监控系统网络安全监测能力验证评估标准", p2_start)
if idx > 0:
    expansions[idx] = [
        "上述建设依据从法律法规、行业标准、企业管理和项目技术四个层面构成了完整的项目实施依据体系。项目实施过程中，我们将严格遵循各层面的要求，确保每一项工作都有据可依、有标准可循。特别是电调〔2023〕28号文件作为项目验收的核心依据，其中明确的各项技术指标和验收标准将作为我们工作质量的最终衡量标准。",
        "在项目合规管理方面，我们建立了与上述依据体系相对应的合规检查机制。在方案设计阶段，核查方案是否满足各项标准要求；在实施阶段，核查操作是否符合规程规定；在验收阶段，核查成果是否达到验收标准。通过全过程的合规管控，确保项目交付成果的合法性和合规性。",
    ]

# === SCHEDULE ADJUSTMENT ===

# Part 1 schedule (paras 73-94): adjust day numbers to fit 2026.7.6-12.31
# Phase 1: 1-20 days -> 2026.7.6-7.25
# Phase 2: 21-40 days -> 2026.7.26-8.14  
# Phase 3: 41-140 days -> 2026.8.15-12.1 (expanded from 41-110)
# Phase 4: 141-180 days -> 2026.12.2-12.31 (expanded from 111-130)

idx = find_para(doc, "第一阶段：充分准备阶段（第1-20天）")
if idx > 0 and idx < 200:
    for run in doc.paragraphs[idx].runs:
        run.text = run.text.replace("第一阶段：充分准备阶段（第1-20天）", "第一阶段：充分准备阶段（2026.7.6-7.25，第1-20天）")

idx = find_para(doc, "第二阶段：分类试点阶段（第21-40天）")
if idx > 0 and idx < 200:
    for run in doc.paragraphs[idx].runs:
        run.text = run.text.replace("第二阶段：分类试点阶段（第21-40天）", "第二阶段：分类试点阶段（2026.7.26-8.14，第21-40天）")

idx = find_para(doc, "第三阶段：全面推进阶段（第41-110天）")
if idx > 0 and idx < 200:
    for run in doc.paragraphs[idx].runs:
        run.text = run.text.replace("第三阶段：全面推进阶段（第41-110天）", "第三阶段：全面推进阶段（2026.8.15-12.1，第41-140天）")

idx = find_para(doc, "第四阶段：验收收尾阶段（第111-130天）")
if idx > 0 and idx < 200:
    for run in doc.paragraphs[idx].runs:
        run.text = run.text.replace("第四阶段：验收收尾阶段（第111-130天）", "第四阶段：验收收尾阶段（2026.12.2-12.31，第141-180天）")

# Part 2 schedule
idx = find_para(doc, "第一阶段：充分准备阶段（第1-15天）")
if idx > 0:
    for run in doc.paragraphs[idx].runs:
        run.text = run.text.replace("第一阶段：充分准备阶段（第1-15天）", "第一阶段：充分准备阶段（2026.7.6-7.20，第1-15天）")

idx = find_para(doc, "第二阶段：试点验证阶段（第16-30天）")
if idx > 0:
    for run in doc.paragraphs[idx].runs:
        run.text = run.text.replace("第二阶段：试点验证阶段（第16-30天）", "第二阶段：试点验证阶段（2026.7.21-8.4，第16-30天）")

idx = find_para(doc, "第三阶段：全面推进阶段（第31-75天）")
if idx > 0:
    for run in doc.paragraphs[idx].runs:
        run.text = run.text.replace("第三阶段：全面推进阶段（第31-75天）", "第三阶段：全面推进阶段（2026.8.5-11.13，第31-130天）")

idx = find_para(doc, "第四阶段：验收收尾阶段（第76-90天）")
if idx > 0:
    for run in doc.paragraphs[idx].runs:
        run.text = run.text.replace("第四阶段：验收收尾阶段（第76-90天）", "第四阶段：验收收尾阶段（2026.11.14-12.31，第131-180天）")

# Apply all text expansions
print(f"\nApplying {len(expansions)} insertion points...")
for para_idx in sorted(expansions.keys(), reverse=True):
    texts = expansions[para_idx]
    if para_idx >= len(doc.paragraphs):
        print(f"  WARNING: {para_idx} out of range")
        continue
    insert_text_after_para(doc, para_idx, texts)
    print(f"  +{len(texts)} after idx {para_idx}")

# Save
print("\nSaving...")
doc.save(INPUT_FILE)
total_chars = sum(len(p.text) for p in doc.paragraphs)
print(f"Final: {len(doc.paragraphs)} paragraphs, {total_chars} chars")
print(f"Estimated pages: ~{total_chars/550 + 2:.0f}")

# Verify hardware sections empty
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == '硬件实力':
        next_t = doc.paragraphs[i+1].text.strip() if i+1 < len(doc.paragraphs) else ''
        print(f"Hardware section at {i}, next='{next_t[:20]}' -> {'OK' if next_t in ['软件实力',''] else 'PROBLEM'}")
        break
