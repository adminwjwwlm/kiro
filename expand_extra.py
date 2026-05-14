# -*- coding: utf-8 -*-
"""Extra expansion to reach ~100 pages - adding substantial content"""
from docx import Document
from docx.shared import Pt
from copy import deepcopy
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INPUT_FILE = '/projects/sandbox/kiro/技术方案/包6-国网新疆昌吉及阿克苏供电公司网络安全项目技术方案.docx'
OUTPUT_FILE = '/projects/sandbox/kiro/技术方案/包6-国网新疆昌吉及阿克苏供电公司网络安全项目技术方案.docx'

def insert_text_after_para(doc, para_idx, texts, font_size=152400):
    ref_para = doc.paragraphs[para_idx]
    for text in reversed(texts):
        new_para = OxmlElement('w:p')
        pPr_source = ref_para._element.find(qn('w:pPr'))
        if pPr_source is not None:
            new_pPr = deepcopy(pPr_source)
            numPr = new_pPr.find(qn('w:numPr'))
            if numPr is not None:
                new_pPr.remove(numPr)
            new_para.append(new_pPr)
        run_elem = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(int(font_size / 12700 * 2)))
        rPr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), str(int(font_size / 12700 * 2)))
        rPr.append(szCs)
        rn = OxmlElement('w:rFonts')
        rn.set(qn('w:eastAsia'), '仿宋_GB2312')
        rPr.append(rn)
        run_elem.append(rPr)
        t = OxmlElement('w:t')
        t.text = text
        t.set(qn('xml:space'), 'preserve')
        run_elem.append(t)
        new_para.append(run_elem)
        ref_para._element.addnext(new_para)

def find_para_index_by_text(doc, search_text, start_from=0):
    for i, para in enumerate(doc.paragraphs):
        if i < start_from:
            continue
        if search_text in para.text:
            return i
    return -1

def find_all_para_by_text(doc, search_text, start_from=0):
    results = []
    for i, para in enumerate(doc.paragraphs):
        if i < start_from:
            continue
        if search_text in para.text:
            results.append(i)
    return results

print("Loading document for extra expansion...")
doc = Document(INPUT_FILE)
print(f"Current: {len(doc.paragraphs)} paragraphs")

expansions = {}
part2_start = find_para_index_by_text(doc, "第二部分 国网新疆阿克苏供电公司")

# === PART 1 - Add more depth to various sections ===

# After "项目目标" section title for Part 1
idx = find_para_index_by_text(doc, "本项目的总体目标是：通过专业、规范、高效的技术服务")
if idx > 0 and idx < part2_start:
    expansions[idx] = [
        "实现上述总体目标需要从技术保障、管理保障和质量保障三个维度同时发力。技术保障方面，需要组建具有丰富升级经验的专业团队，配备完善的工具设备，制定科学严谨的操作方案。管理保障方面，需要建立高效的项目管理机制，确保各项工作按计划有序推进。质量保障方面，需要建立多层次的质量检查机制，确保每一台设备的升级质量达到标准要求。",
        "为确保总体目标的顺利实现，我们将项目管理贯穿于项目全生命周期，从项目启动、计划制定、实施执行、监控调整到收尾验收，每个阶段都设有明确的工作内容、交付成果和质量标准。项目经理全面负责项目的进度、质量和风险管理，确保各项工作有序协调推进。",
    ]

# Expand the 单次升级详细计划 description
idx = find_para_index_by_text(doc, "每次全面升级的详细工作计划如下")
if idx > 0 and idx < part2_start:
    expansions[idx] = [
        "单次全面升级按照\"准备-执行-验证-归档\"四个阶段组织实施，总周期约为2-3周（含前期准备和后期归档）。其中现场升级执行阶段约需8-10个工作日完成全部19座变电站的升级工作。详细的日程安排如下表所示，表中展示了标准情况下的升级批次安排。实际执行时可能根据天气、交通、变电站运行情况等因素进行适当调整。",
        "在人员安排方面，每次全面升级配置2个升级小组（升级组A和升级组B），每组由2名技术人员组成（1名主操作手+1名辅助操作和记录人员）。两个小组可以同时在不同变电站开展升级工作，在确保每组操作质量的同时提高整体升级效率。项目管理组由项目经理和质量管控人员组成，负责统筹协调和质量监督。",
    ]

# Expand 沟通协调计划
idx = find_para_index_by_text(doc, "沟通协调计划")
if idx > 0 and idx < part2_start:
    expansions[idx + 4] = [
        "（5）问题升级机制：对于升级过程中遇到的无法在现场立即解决的技术问题，按照问题严重程度分级上报。一般问题由项目经理协调解决；重大问题由双方技术负责人会商解决；紧急问题启动应急响应机制处理。",
        "（6）信息记录制度：所有重要的沟通协调信息均做好书面记录，包括沟通时间、参与人员、沟通内容、达成共识和待办事项等，确保信息传递的准确性和可追溯性。",
        "项目质量管控计划",
        "为确保项目实施质量，建立以下质量管控措施：",
        "（1）操作前检查：每次升级操作前，由操作人员自检确认各项准备工作就绪，包括升级包完整性、工具设备齐备、操作方案熟悉等。自检通过后开始操作。",
        "（2）操作中监控：升级操作过程中，辅助人员全程监控设备状态和操作步骤，及时发现和纠正任何偏差。关键步骤需要双人确认后才能执行。",
        "（3）操作后验证：每台设备升级完成后，按照验证检查表逐项进行验证，所有项目通过后才确认该设备升级成功。验证结果由操作人员和验证人员共同签字确认。",
        "（4）日报审核：项目管理组每天审核各升级小组提交的工作日报和升级记录，发现质量问题及时反馈纠正。",
        "（5）阶段评审：每个升级批次完成后，组织阶段评审会议，总结经验教训，改进后续批次的操作质量。",
    ]

# === PART 2 - Add more depth ===

# Expand Part 2 项目理解 further
idx = find_para_index_by_text(doc, "项目实施遵循", part2_start)
if idx > 0:
    expansions[idx] = [
        "从项目管理角度看，本项目具有以下特点需要重点关注：一是工作内容多样，涉及设备接入、策略调优、安全加固、漏洞修复、辨识提升、端口封堵等六大类工作，各类工作的技术要求和操作方式各不相同，需要不同专业背景的技术人员协同配合；二是实施环境敏感，所有工作在运行中的变电站内进行，必须确保任何操作不影响电力系统的正常运行；三是成果可量化，各项工作都有明确的量化目标和验证标准，项目效果可以通过数据进行客观评估。",
        "基于上述特点，我们采用\"统一规划、分工协作、风险可控、效果可验\"的项目管理策略。统一规划确保各项工作有序衔接不冲突；分工协作确保各专业人员高效配合；风险可控确保每一步操作都在安全范围内；效果可验确保项目成果经得起检验。",
    ]

# Expand Part 2 设备接入 more detail
idx = find_para_index_by_text(doc, "设备接入台账", part2_start)
if idx > 0:
    expansions[idx + 3] = [
        "设备接入工作的实施步骤详细描述如下：",
        "步骤一：现场设备核查。技术人员进入变电站后，首先对照已有资料逐一核实站内设备的实际情况，包括设备的物理位置、型号铭牌信息、网络连接方式等。对于资料中未记录的设备或信息有出入的设备，现场补充完善台账信息。",
        "步骤二：接入条件评估。对每台待接入设备评估其接入条件，包括：设备是否支持Syslog日志输出、是否支持SNMP协议、是否已安装安全Agent、网络连通性是否满足数据传输要求等。根据评估结果确定每台设备的最优接入方式。",
        "步骤三：配置实施。根据确定的接入方式，在设备侧和监测系统侧分别进行配置。设备侧配置主要包括：启用日志输出功能、配置日志输出目标地址和端口、设置日志输出级别和内容范围等。监测系统侧配置主要包括：添加数据源、配置解析规则、设置告警关联条件等。",
        "步骤四：接入验证。配置完成后进行接入效果验证：在设备上触发一个测试事件（如登录失败），确认该事件能在安全监测系统中被正确接收、解析和展示。验证成功后该设备接入工作完成。",
    ]

# Expand Part 2 漏洞扫描 more
idx = find_para_index_by_text(doc, "漏洞扫描报告", part2_start)
if idx > 0:
    expansions[idx + 3] = [
        "漏洞扫描与修复工作的技术要点补充说明：",
        "（1）扫描时间选择：漏洞扫描操作安排在设备负荷较低的时段进行，避免扫描流量对设备正常运行产生影响。对于扫描可能产生较大网络流量的情况，提前与运维人员沟通并做好监控准备。",
        "（2）扫描策略配置：根据电力监控系统的特殊性，配置适合的扫描策略。避免使用侵入性强的扫描方式（如暴力破解测试），防止扫描操作本身对设备造成影响。对于敏感设备，采用被动扫描或降低扫描强度。",
        "（3）漏洞去重和关联分析：对扫描结果进行去重处理，将同一漏洞在不同设备上的出现情况进行汇总分析；识别漏洞之间的关联关系，评估漏洞组合利用的风险等级。",
        "（4）修复影响预评估：对每一个计划修复的漏洞，在实施修复之前进行影响预评估。评估内容包括：修复操作是否需要重启设备、是否影响设备功能、是否有兼容性风险、修复失败的后果等。只有评估结果可接受的漏洞才进入实施修复环节。",
    ]

# Expand Part 2 工作计划 - 沟通协调
idx = find_para_index_by_text(doc, "验收评审会", part2_start)
if idx > 0:
    expansions[idx] = [
        "项目质量管控计划",
        "本项目建立三级质量管控体系，确保各项工作质量：",
        "（1）操作级质量控制：每项操作执行前进行方案复核，执行中实行双人复核制度，执行后进行即时效果验证。确保每一步操作都准确无误。",
        "（2）模块级质量控制：每完成一个工作模块（如设备接入、漏洞修复等），进行模块级的质量评审，核查该模块工作的完整性和有效性，确认达到预定目标。",
        "（3）项目级质量控制：全部工作完成后进行项目级的综合质量评审，全面核查各项目标指标的达成情况，编制项目质量报告。",
        "人员安全管理计划",
        "项目实施期间的人员安全管理措施包括：",
        "（1）所有进站人员必须经过变电站安全培训并取得进站作业资格。",
        "（2）现场作业严格遵守变电站安全规程，佩戴必要的安全防护用品。",
        "（3）涉及设备操作的工作必须在站端运维人员监护下进行。",
        "（4）项目团队人员签订保密协议，不得将在工作中获取的设备配置信息、网络架构信息等敏感信息向外泄露。",
        "（5）使用的便携设备和存储介质经过严格的安全检测，进站前查杀病毒，出站时清除敏感数据。",
    ]

# Apply all expansions
print(f"\nTotal extra insertion points: {len(expansions)}")
for para_idx in sorted(expansions.keys(), reverse=True):
    texts = expansions[para_idx]
    if para_idx >= len(doc.paragraphs):
        print(f"  WARNING: para_idx {para_idx} out of range, skipping")
        continue
    insert_text_after_para(doc, para_idx, texts)
    print(f"  Inserted {len(texts)} paragraphs after index {para_idx}")

print(f"\nTotal paragraphs: {len(doc.paragraphs)}")
print("Saving...")
doc.save(OUTPUT_FILE)

total_chars = sum(len(p.text) for p in doc.paragraphs)
est_pages = total_chars / 800
print(f"Total characters: {total_chars}")
print(f"Estimated pages: {est_pages:.0f}")
